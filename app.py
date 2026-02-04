import streamlit as st
import json
import re
import os
import hashlib
from pathlib import Path
from groq import Groq
import PyPDF2
import io
import pandas as pd
from dotenv import load_dotenv
from docx import Document
from fpdf import FPDF
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from rapidfuzz import fuzz, process

# Load environment variables from .env file
load_dotenv()

# Get API key from environment
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Page config
st.set_page_config(
    page_title="HireSmartAI - Recruitment OS",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for modern UI with purple theme
st.markdown("""
<style>
    /* Hide Streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* Main background */
    .stApp {
        background: #f8f9fc;
    }

    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background: white;
        border-right: 1px solid #e5e7eb;
        padding-top: 1rem;
    }

    [data-testid="stSidebar"] .stMarkdown {
        padding: 0 1rem;
    }

    /* Logo and brand */
    .brand-container {
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 1rem;
        margin-bottom: 1rem;
    }

    .brand-icon {
        width: 40px;
        height: 40px;
        background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
        border-radius: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-size: 20px;
    }

    .brand-text {
        font-weight: 700;
        font-size: 1.1rem;
        color: #1f2937;
    }

    .brand-subtitle {
        font-size: 0.75rem;
        color: #6b7280;
    }

    /* Navigation items */
    .nav-item {
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 0.75rem 1rem;
        border-radius: 8px;
        margin: 0.25rem 0;
        cursor: pointer;
        transition: all 0.2s;
        color: #4b5563;
        text-decoration: none;
    }

    .nav-item:hover {
        background: #f3f4f6;
    }

    .nav-item.active {
        background: linear-gradient(135deg, #ede9fe 0%, #ddd6fe 100%);
        color: #7c3aed;
        font-weight: 600;
    }

    /* Main content area */
    .main-header {
        font-size: 1.75rem;
        font-weight: 700;
        color: #1f2937;
        margin-bottom: 0.25rem;
    }

    .main-subtitle {
        color: #6b7280;
        font-size: 0.95rem;
        margin-bottom: 1.5rem;
    }

    /* Cards */
    .upload-card, .jd-card {
        background: white;
        border-radius: 16px;
        padding: 1.5rem;
        border: 2px dashed #e5e7eb;
        min-height: 200px;
    }

    .jd-card {
        border: 1px solid #e5e7eb;
        border-radius: 16px;
    }

    /* File chip */
    .file-chip {
        display: inline-flex;
        align-items: center;
        gap: 8px;
        background: #f9fafb;
        border: 1px solid #e5e7eb;
        border-radius: 8px;
        padding: 0.5rem 0.75rem;
        margin: 0.25rem;
        font-size: 0.85rem;
        color: #374151;
    }

    /* Results section */
    .results-header {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin: 2rem 0 1rem 0;
    }

    .results-title {
        font-size: 1.25rem;
        font-weight: 600;
        color: #1f2937;
        display: flex;
        align-items: center;
        gap: 12px;
    }

    .match-count {
        background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
    }

    /* Category section headers */
    .category-header {
        display: flex;
        align-items: center;
        gap: 12px;
        margin: 1.5rem 0 1rem 0;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #e5e7eb;
    }

    .category-title {
        font-size: 1.1rem;
        font-weight: 600;
        color: #1f2937;
    }

    .category-count {
        padding: 0.2rem 0.6rem;
        border-radius: 12px;
        font-size: 0.8rem;
        font-weight: 600;
    }

    /* Skill tags */
    .skill-matched {
        display: inline-block;
        background: #dcfce7;
        color: #166534;
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        font-size: 0.75rem;
        margin: 0.1rem;
        border: 1px solid #bbf7d0;
    }

    .skill-missing {
        display: inline-block;
        background: #fee2e2;
        color: #991b1b;
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        font-size: 0.75rem;
        margin: 0.1rem;
        border: 1px solid #fecaca;
    }

    .skill-nice {
        display: inline-block;
        background: #e0e7ff;
        color: #3730a3;
        padding: 0.25rem 0.5rem;
        border-radius: 4px;
        font-size: 0.75rem;
        margin: 0.1rem;
        border: 1px solid #c7d2fe;
    }

    /* Section labels */
    .section-label {
        font-size: 0.75rem;
        font-weight: 600;
        color: #9ca3af;
        text-transform: uppercase;
        letter-spacing: 0.05em;
        margin-bottom: 0.5rem;
    }

    /* Custom streamlit overrides */
    .stTextArea textarea {
        border-radius: 12px;
        border: 1px solid #e5e7eb;
        padding: 1rem;
        font-size: 0.95rem;
    }

    .stTextArea textarea:focus {
        border-color: #8b5cf6;
        box-shadow: 0 0 0 3px rgba(139, 92, 246, 0.1);
    }

    div[data-testid="stFileUploader"] {
        background: white;
        border-radius: 16px;
        padding: 1rem;
        border: 2px dashed #ddd6fe;
    }

    div[data-testid="stFileUploader"]:hover {
        border-color: #8b5cf6;
    }

    .stButton > button {
        background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 12px;
        font-weight: 600;
        font-size: 1rem;
        transition: all 0.2s;
        box-shadow: 0 4px 14px rgba(124, 58, 237, 0.25);
    }

    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(124, 58, 237, 0.35);
        background: linear-gradient(135deg, #7c3aed 0%, #6d28d9 100%);
    }

    .stProgress > div > div {
        background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
    }

    /* Keyword tags */
    .keyword-tag {
        display: inline-block;
        background: #ede9fe;
        color: #7c3aed;
        padding: 0.3rem 0.6rem;
        border-radius: 6px;
        font-size: 0.8rem;
        margin: 0.2rem;
        border: 1px solid #ddd6fe;
    }

    /* User profile in sidebar */
    .user-profile {
        display: flex;
        align-items: center;
        gap: 12px;
        padding: 1rem;
        border-top: 1px solid #e5e7eb;
        margin-top: auto;
    }

    .user-avatar {
        width: 40px;
        height: 40px;
        border-radius: 50%;
        background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-weight: 600;
    }

    .user-name {
        font-weight: 600;
        color: #1f2937;
        font-size: 0.9rem;
    }

    .user-plan {
        font-size: 0.75rem;
        color: #8b5cf6;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'results' not in st.session_state:
    st.session_state.results = []
if 'analyzed' not in st.session_state:
    st.session_state.analyzed = False
if 'folder_pdf_paths' not in st.session_state:
    st.session_state.folder_pdf_paths = []
if 'duplicates_count' not in st.session_state:
    st.session_state.duplicates_count = 0
if 'total_files_processed' not in st.session_state:
    st.session_state.total_files_processed = 0
if 'suggested_keywords' not in st.session_state:
    st.session_state.suggested_keywords = []
if 'job_title' not in st.session_state:
    st.session_state.job_title = ""

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_pdf_path(pdf_path):
    """Extract text from PDF file path (for folder upload)"""
    try:
        with open(pdf_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            text = re.sub(r'\s+', ' ', text).strip()
            return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_docx(docx_file):
    """Extract text from uploaded DOCX file"""
    try:
        doc = Document(io.BytesIO(docx_file.read()))
        text = ""
        for para in doc.paragraphs:
            text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_docx_path(docx_path):
    """Extract text from DOCX file path (for folder upload)"""
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + " "
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def get_content_hash(text):
    """Generate hash of resume content for duplicate detection"""
    normalized = re.sub(r'\s+', ' ', text.lower().strip())
    return hashlib.md5(normalized.encode()).hexdigest()

def extract_keywords_from_jd(client, job_description):
    """Extract suggested keywords from job description using AI"""
    prompt = f"""Analyze the following job description and extract key skills, technologies, and qualifications.

Return ONLY valid JSON in this exact format:
{{
    "job_title": "<extracted job title>",
    "must_have_skills": ["skill1", "skill2", "skill3"],
    "nice_to_have_skills": ["skill1", "skill2"],
    "technologies": ["tech1", "tech2"],
    "qualifications": ["qual1", "qual2"],
    "experience_required": "<e.g., 3-5 years>"
}}

JOB DESCRIPTION:
{job_description[:3000]}

Return ONLY the JSON object."""

    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=800,
            seed=42
        )
        result_text = response.choices[0].message.content.strip()

        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0]
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0]

        return json.loads(result_text)
    except:
        return {"job_title": "", "must_have_skills": [], "nice_to_have_skills": [], "technologies": [], "qualifications": [], "experience_required": ""}

## ===================== ML SCORING FUNCTIONS ===================== ##

def extract_contact_info(text):
    """Extract email, phone, location using regex - NO LLM needed"""
    # Email
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    emails = re.findall(email_pattern, text)
    email = emails[0] if emails else "Not provided"

    # Phone (Indian + international formats)
    phone_patterns = [
        r'(?:\+91[\s-]?)?[6-9]\d{4}[\s-]?\d{5}',  # Indian mobile
        r'(?:\+91[\s-]?)?\d{5}[\s-]?\d{5}',          # 10 digit
        r'(?:\+\d{1,3}[\s-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}',  # US/intl
        r'\d{3}[\s-]\d{3}[\s-]\d{4}',                 # XXX-XXX-XXXX
    ]
    phone = "Not provided"
    for pattern in phone_patterns:
        phones = re.findall(pattern, text)
        if phones:
            phone = phones[0].strip()
            break

    # Location - common Indian cities + international
    cities = [
        "Mumbai", "Delhi", "Bangalore", "Bengaluru", "Hyderabad", "Chennai",
        "Kolkata", "Pune", "Ahmedabad", "Jaipur", "Lucknow", "Kanpur",
        "Nagpur", "Indore", "Thane", "Bhopal", "Visakhapatnam", "Vadodara",
        "Gurgaon", "Gurugram", "Noida", "Chandigarh", "Coimbatore", "Kochi",
        "Mysore", "Mysuru", "Surat", "Nashik", "Rajkot", "Ranchi",
        "New York", "San Francisco", "London", "Dubai", "Singapore", "Toronto",
        "Remote", "Work from home", "WFH", "Hybrid"
    ]
    location = "Not provided"
    text_lower = text.lower()
    for city in cities:
        if city.lower() in text_lower:
            location = city
            break

    return email, phone, location

def extract_experience_years(text):
    """Extract years of experience using regex - NO LLM needed"""
    patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
        r'experience\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)',
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:in|of)\s*(?:IT|software|development|engineering)',
        r'total\s*(?:experience|exp)\s*(?:of)?\s*(\d+)',
        r'(\d+)\s*(?:years?|yrs?)\s*(\d+)\s*(?:months?|mos?)',
    ]
    years = 0
    for pattern in patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            match = matches[0]
            if isinstance(match, tuple):
                years = int(match[0])
            else:
                years = int(match)
            break
    return years

def extract_required_experience(jd_text):
    """Extract required years from JD"""
    patterns = [
        r'(\d+)\s*[-–to]+\s*(\d+)\+?\s*(?:years?|yrs?)',  # "5-10+ years", "5-10 years"
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',  # "5+ years of experience"
        r'minimum\s*(\d+)\s*(?:years?|yrs?)',  # "minimum 5 years"
        r'(\d+)\+?\s*(?:years?|yrs?)',  # Simple fallback: "5+ years", "5 years"
    ]
    for pattern in patterns:
        matches = re.findall(pattern, jd_text, re.IGNORECASE)
        if matches:
            match = matches[0]
            if isinstance(match, tuple):
                return int(match[0])  # Take lower bound
            return int(match)
    return 0

def extract_education(text):
    """Detect education level from resume text - NO LLM needed"""
    text_lower = text.lower()

    # Check for degrees (highest first)
    phd_keywords = ['ph.d', 'phd', 'doctorate', 'doctoral']
    masters_keywords = ['m.tech', 'mtech', 'm.sc', 'msc', 'mba', 'm.e.', 'masters', 'master of', 'ms in', 'm.s.', 'mca', 'm.c.a']
    bachelors_keywords = ['b.tech', 'btech', 'b.sc', 'bsc', 'b.e.', 'bachelor', 'bca', 'b.c.a', 'b.eng', 'beng', 'b.com']
    diploma_keywords = ['diploma', 'polytechnic', 'certification', 'certified']

    for kw in phd_keywords:
        if kw in text_lower:
            return "PhD", 10
    for kw in masters_keywords:
        if kw in text_lower:
            return "Masters", 10
    for kw in bachelors_keywords:
        if kw in text_lower:
            return "Bachelors", 7
    for kw in diploma_keywords:
        if kw in text_lower:
            return "Diploma", 5

    return "Not detected", 3

def extract_skills_from_jd(jd_text):
    """Extract required skills from job description using keyword patterns"""
    # Common tech skills to look for
    common_skills = [
        "Python", "Java", "JavaScript", "TypeScript", "C++", "C#", "Go", "Rust", "Ruby",
        "React", "Angular", "Vue", "Node.js", "Express", "Django", "Flask", "FastAPI",
        "Spring Boot", "Spring", ".NET", "Laravel", "Rails",
        "SQL", "MySQL", "PostgreSQL", "MongoDB", "Redis", "Elasticsearch", "Cassandra",
        "AWS", "Azure", "GCP", "Docker", "Kubernetes", "Terraform", "Jenkins", "CI/CD",
        "Git", "GitHub", "GitLab", "Bitbucket",
        "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "AI",
        "TensorFlow", "PyTorch", "Scikit-learn", "Pandas", "NumPy",
        "Power BI", "Tableau", "Excel", "Data Analysis", "Data Science", "Data Engineering",
        "REST", "API", "GraphQL", "Microservices", "SOA",
        "Linux", "Unix", "Windows Server", "Networking",
        "Agile", "Scrum", "JIRA", "Confluence",
        "HTML", "CSS", "SASS", "Bootstrap", "Tailwind",
        "Spark", "Hadoop", "Kafka", "Airflow", "ETL",
        "Selenium", "JUnit", "pytest", "Testing", "QA",
        "Figma", "Photoshop", "UI/UX",
        "SAP", "Salesforce", "ServiceNow", "Oracle",
        "Pyramid Analytics", "model monitoring", "drift detection",
        "R", "SAS", "SPSS", "MATLAB",
        "communication", "leadership", "problem-solving", "teamwork",
    ]

    jd_lower = jd_text.lower()
    found_skills = []
    for skill in common_skills:
        if skill.lower() in jd_lower:
            found_skills.append(skill)

    # Also extract quoted or bulleted skills
    bullet_pattern = r'[•\-\*]\s*([A-Za-z][A-Za-z\s/\.#\+]{2,30})'
    bullets = re.findall(bullet_pattern, jd_text)
    for b in bullets:
        b_clean = b.strip()
        if b_clean and b_clean not in found_skills and len(b_clean) < 30:
            found_skills.append(b_clean)

    return sorted(set(found_skills), key=str.lower) if found_skills else ["General skills"]

def fuzzy_match_skills(required_skills, resume_text, threshold=70):
    """Match skills using fuzzy string matching - handles typos and variations"""
    resume_lower = resume_text.lower()
    matched = []
    missing = []

    for skill in required_skills:
        skill_lower = skill.lower()

        # Exact match first
        if skill_lower in resume_lower:
            matched.append(skill)
            continue

        # Fuzzy match - check against words/phrases in resume
        resume_words = re.findall(r'[a-zA-Z][a-zA-Z\s/\.#\+]{2,30}', resume_text)
        best_match = process.extractOne(skill_lower, [w.lower() for w in resume_words], scorer=fuzz.ratio)

        if best_match and best_match[1] >= threshold:
            matched.append(skill)
        else:
            missing.append(skill)

    return matched, missing

def fuzzy_match_nice_to_have(nice_to_have_skills, resume_text, threshold=70):
    """Match nice-to-have skills using fuzzy matching"""
    if not nice_to_have_skills:
        return []
    resume_lower = resume_text.lower()
    matched = []
    for skill in nice_to_have_skills:
        skill_lower = skill.lower()
        if skill_lower in resume_lower:
            matched.append(skill)
            continue
        resume_words = re.findall(r'[a-zA-Z][a-zA-Z\s/\.#\+]{2,30}', resume_text)
        best_match = process.extractOne(skill_lower, [w.lower() for w in resume_words], scorer=fuzz.ratio)
        if best_match and best_match[1] >= threshold:
            matched.append(skill)
    return matched

def calculate_tfidf_similarity(resume_text, jd_text):
    """Calculate text similarity between resume and JD using TF-IDF"""
    try:
        vectorizer = TfidfVectorizer(stop_words='english', max_features=5000)
        tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
        return round(similarity * 100, 2)
    except:
        return 0

def calculate_verdict_from_score(score):
    """Calculate verdict deterministically from score - NO LLM involvement"""
    score = max(0, min(100, score))

    if score >= 85:
        return "Best Fit", "Recommend for Interview"
    elif score >= 70:
        return "Strong Fit", "Recommend for Interview"
    elif score >= 50:
        return "Average", "Consider for Interview"
    else:
        return "Not a Fit", "Do Not Recommend"


## ===================== NAME EXTRACTION ===================== ##

def split_camel_case(text):
    """Split CamelCase into separate words: 'AnkitDarade' -> 'Ankit Darade'"""
    result = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    result = re.sub(r'([A-Z]+)([A-Z][a-z])', r'\1 \2', result)
    return result

def clean_candidate_name(file_name):
    """Remove platform prefixes, experience suffixes, and clean up filename for display"""
    name = file_name

    # Remove file extensions
    for ext in ['.pdf', '.PDF', '.docx', '.DOCX', '.doc', '.DOC']:
        name = name.replace(ext, '')

    # Remove common job portal prefixes (case-insensitive)
    platform_prefixes = [
        r'(?i)^naukri[_\-\s]+',
        r'(?i)^indeed[_\-\s]+',
        r'(?i)^linkedin[_\-\s]+',
        r'(?i)^monster[_\-\s]+',
        r'(?i)^shine[_\-\s]+',
        r'(?i)^timesjobs[_\-\s]+',
        r'(?i)^glassdoor[_\-\s]+',
        r'(?i)^foundit[_\-\s]+',
        r'(?i)^ziprecruiter[_\-\s]+',
        r'(?i)^hirect[_\-\s]+',
        r'(?i)^instahyre[_\-\s]+',
        r'(?i)^apna[_\-\s]+',
        r'(?i)^iimjobs[_\-\s]+',
        r'(?i)^cutshort[_\-\s]+',
        r'(?i)^hirist[_\-\s]+',
        r'(?i)^angellist[_\-\s]+',
        r'(?i)^wellfound[_\-\s]+',
    ]
    for prefix in platform_prefixes:
        name = re.sub(prefix, '', name)

    # Remove experience tags like [6y_0m], [5y 6m], (6y_0m), [2y_3m], etc.
    name = re.sub(r'[\[\(]\d+y[\s_\-]?\d*m?[\]\)]', '', name)

    # Remove trailing numbers in brackets/parens like (1), [1], (2), [2]
    name = re.sub(r'[\[\(]\d+[\]\)]', '', name)

    # Remove date patterns like _12_01_26, _12-01-2026, _2026_01_12
    name = re.sub(r'[_\-]?\d{1,2}[_\-]\d{1,2}[_\-]\d{2,4}', '', name)
    name = re.sub(r'[_\-]?\d{4}[_\-]\d{1,2}[_\-]\d{1,2}', '', name)

    # Replace underscores and hyphens with spaces
    name = name.replace('_', ' ').replace('-', ' ')

    # Split CamelCase: "AnkitDarade" -> "Ankit Darade"
    name = split_camel_case(name)

    # Remove "Resume", "CV", etc. AFTER underscores are replaced (so \b works correctly)
    name = re.sub(r'(?i)\b(resume|cv|curriculum\s*vitae)\b', '', name)

    # Remove extra whitespace
    name = re.sub(r'\s+', ' ', name).strip()

    # Remove any remaining lone numbers
    name = re.sub(r'\b\d+\b', '', name).strip()
    name = re.sub(r'\s+', ' ', name).strip()

    return name if name else file_name


def extract_name_from_resume(resume_text):
    """Extract candidate name from resume text content.
    Looks at the first few lines for a name pattern (2-4 capitalized words)."""
    if not resume_text or resume_text.startswith("Error"):
        return None

    # Words that are NOT names - common resume headers/labels
    skip_words = {
        'resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary',
        'objective', 'experience', 'education', 'skills', 'contact',
        'phone', 'email', 'address', 'personal', 'details', 'information',
        'professional', 'career', 'about', 'me', 'page', 'confidential',
        'private', 'strictly', 'naukri', 'indeed', 'linkedin', 'http',
        'https', 'www', 'gmail', 'yahoo', 'hotmail', 'outlook',
        'mobile', 'tel', 'name', 'date', 'gender', 'nationality',
        'references', 'available', 'upon', 'request', 'dear', 'sir',
        'madam', 'to', 'whom', 'it', 'may', 'concern'
    }

    # Get first 15 non-empty lines of the resume
    lines = resume_text.strip().split('\n')
    candidate_lines = []
    for line in lines:
        cleaned = line.strip()
        if cleaned and len(cleaned) > 1:
            candidate_lines.append(cleaned)
        if len(candidate_lines) >= 15:
            break

    for line in candidate_lines:
        # Skip lines that are too long (likely paragraphs, not names)
        if len(line) > 50:
            continue

        # Skip lines with email addresses
        if '@' in line:
            continue

        # Skip lines with 7+ consecutive digits (phone numbers)
        if re.search(r'\d{7,}', line):
            continue

        # Skip lines that are mostly numbers
        digits = sum(c.isdigit() for c in line)
        if digits > len(line) * 0.3:
            continue

        # Skip lines with URLs
        if re.search(r'https?://|www\.', line, re.IGNORECASE):
            continue

        # Skip lines that are common headers
        line_lower = line.lower().strip()
        if line_lower in skip_words:
            continue
        if any(line_lower.startswith(w) for w in ['objective', 'summary', 'experience', 'education', 'skills', 'profile summary', 'contact', 'phone', 'email', 'mobile', 'address']):
            continue

        # Clean the line - remove special chars at edges
        clean_line = re.sub(r'^[^a-zA-Z]+|[^a-zA-Z]+$', '', line.strip())
        if not clean_line:
            continue

        # Try to match a name pattern: 2-4 words, predominantly letters
        words = clean_line.split()
        if 2 <= len(words) <= 4 and len(clean_line) < 40:
            all_name_like = all(
                len(w) >= 1 and w[0].isupper() and re.match(r'^[A-Za-z.]+$', w)
                for w in words
            )
            if all_name_like:
                lower_words = [w.lower().rstrip('.') for w in words]
                if not any(w in skip_words for w in lower_words):
                    return clean_line

        # Also try: single CamelCase word that splits into 2-3 name parts
        if len(words) == 1 and len(clean_line) >= 4:
            split_name = split_camel_case(clean_line)
            split_words = split_name.split()
            if 2 <= len(split_words) <= 3:
                all_alpha = all(w.isalpha() and w[0].isupper() for w in split_words)
                if all_alpha:
                    lower_words = [w.lower() for w in split_words]
                    if not any(w in skip_words for w in lower_words):
                        return split_name

    return None


## ===================== HYBRID ANALYSIS ===================== ##

def analyze_resume(client, resume_text, job_description, nice_to_have_skills, candidate_name, job_title):
    """HYBRID ML + LLM resume analysis - ML for scoring, LLM for summary only"""

    # ==================== STEP 1: ML-BASED SCORING (Deterministic) ====================

    # 1a. Extract required skills from JD
    required_skills = extract_skills_from_jd(job_description)

    # 1b. Skills matching with fuzzy logic (0-40 points)
    skills_matched, skills_missing = fuzzy_match_skills(required_skills, resume_text)
    if required_skills:
        skills_score = round((len(skills_matched) / len(required_skills)) * 40)
    else:
        skills_score = 20  # Default if no skills detected
    skills_score = min(40, skills_score)

    # 1c. Experience extraction and scoring (0-25 points)
    candidate_years = extract_experience_years(resume_text)
    required_years = extract_required_experience(job_description)
    if required_years > 0:
        exp_ratio = min(candidate_years / required_years, 1.5)  # Cap at 150%
        experience_score = round(exp_ratio * 25)
    elif candidate_years > 0:
        experience_score = min(25, candidate_years * 3)  # 3 points per year
    else:
        experience_score = 10  # Default
    experience_score = min(25, experience_score)

    # 1d. Nice-to-have skills matching (0-15 points)
    nice_matched = fuzzy_match_nice_to_have(nice_to_have_skills, resume_text)
    nice_to_have_score = min(15, len(nice_matched) * 5)

    # 1e. Education detection (0-10 points)
    education_level, education_score = extract_education(resume_text)

    # 1f. TF-IDF relevance score (0-10 points)
    tfidf_sim = calculate_tfidf_similarity(resume_text, job_description)
    relevance_score = min(10, round(tfidf_sim / 10))  # Convert 0-100 to 0-10

    # ==================== TOTAL ML SCORE ====================
    total_score = skills_score + experience_score + nice_to_have_score + education_score + relevance_score
    total_score = min(100, max(0, total_score))

    # Deterministic verdict
    verdict, recommendation = calculate_verdict_from_score(total_score)

    # Contact info extraction (regex)
    email, phone, location = extract_contact_info(resume_text)

    # ==================== JD DUPLICATE DETECTION ====================
    is_jd_duplicate = tfidf_sim > 95  # Flag if resume is >95% similar to JD

    # ==================== STEP 2: LLM FOR SUMMARY + NAME ====================
    summary = ""
    current_role = "Not specified"
    strengths = []
    weaknesses = []
    llm_candidate_name = ""

    try:
        prompt = f"""Analyze this resume briefly. Return ONLY valid JSON:
{{
    "candidate_name": "<full name of the candidate from the resume>",
    "current_role": "<current/latest job title>",
    "strengths": ["strength1", "strength2"],
    "weaknesses": ["weakness1"],
    "summary": "<2 sentence evaluation of this candidate for {job_title} role>"
}}

RESUME (first 2000 chars):
{resume_text[:2000]}

Return ONLY the JSON object."""

        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=400,
            seed=42
        )
        result_text = response.choices[0].message.content.strip()

        if "```json" in result_text:
            result_text = result_text.split("```json")[1].split("```")[0]
        elif "```" in result_text:
            result_text = result_text.split("```")[1].split("```")[0]

        llm_result = json.loads(result_text)
        llm_candidate_name = llm_result.get("candidate_name", "")
        current_role = llm_result.get("current_role", "Not specified")
        strengths = llm_result.get("strengths", [])
        weaknesses = llm_result.get("weaknesses", [])
        summary = llm_result.get("summary", "")

    except:
        summary = f"Candidate has {candidate_years} years experience. Matched {len(skills_matched)}/{len(required_skills)} required skills."
        current_role = "Not specified"
        strengths = [f"{len(skills_matched)} skills matched"] if skills_matched else []
        weaknesses = [f"{len(skills_missing)} skills missing"] if skills_missing else []

    # Use LLM-extracted name if the current name looks like a job title or generic text
    final_name = candidate_name
    if llm_candidate_name and len(llm_candidate_name) > 2:
        # Check if current name looks like a real person name (2-4 words, all alpha)
        current_words = candidate_name.split()
        looks_like_name = (
            2 <= len(current_words) <= 4 and
            all(w.isalpha() for w in current_words) and
            len(candidate_name) < 40
        )
        if not looks_like_name:
            # Current name doesn't look like a person - use LLM name
            final_name = llm_candidate_name
        elif any(kw in candidate_name.lower() for kw in ['senior', 'junior', 'manager', 'engineer', 'developer', 'analyst', 'scientist', 'lead', 'director', 'consultant', 'intern', 'associate', 'executive']):
            # Current name contains job title keywords - use LLM name
            final_name = llm_candidate_name

    # Override for JD duplicates
    if is_jd_duplicate:
        verdict = "Not a Fit"
        recommendation = "Possible JD Upload - Not a Resume"
        summary = "WARNING: This file appears to be the Job Description itself (99%+ similarity), not a candidate resume."

    # ==================== RETURN COMBINED RESULT ====================
    return {
        "candidate_name": final_name,
        "job_title": job_title,
        "fit_score": total_score,
        "verdict": verdict,
        "recommendation": recommendation,
        "email": email,
        "phone": phone,
        "location": location,
        "current_role": current_role,
        "experience_years": candidate_years,
        "education_level": education_level,
        "skills_matched": skills_matched,
        "skills_missing": skills_missing,
        "nice_to_have_matched": nice_matched,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "summary": summary,
        "tfidf_similarity": tfidf_sim,
        "score_breakdown": {
            "skills_score": skills_score,
            "experience_score": experience_score,
            "nice_to_have_score": nice_to_have_score,
            "education_score": education_score,
            "relevance_score": relevance_score
        }
    }

def get_category_color(category):
    colors = {
        "Best Fit": ("#7c3aed", "#ede9fe"),
        "Strong Fit": ("#059669", "#d1fae5"),
        "Average": ("#d97706", "#fef3c7"),
        "Not a Fit": ("#dc2626", "#fee2e2")
    }
    return colors.get(category, ("#6b7280", "#f3f4f6"))


def create_excel_report(results, categories, total_files, duplicates, job_title):
    """Create Excel report with all candidate data"""
    data = []
    for result in results:
        data.append({
            "Candidate Name": result.get('candidate_name', 'Unknown'),
            "Fit Score": result.get('fit_score', 0),
            "Verdict": result.get('verdict', 'N/A'),
            "Job Title Applied": job_title,
            "Current Role": result.get('current_role', 'N/A'),
            "Experience (Years)": result.get('experience_years', 0),
            "Location": result.get('location', 'Not provided'),
            "Email": result.get('email', 'Not provided'),
            "Phone": result.get('phone', 'Not provided'),
            "Skills Matched": ", ".join(result.get('skills_matched', [])),
            "Skills Missing": ", ".join(result.get('skills_missing', [])),
            "Nice-to-Have Skills": ", ".join(result.get('nice_to_have_matched', [])),
            "Strengths": ", ".join(result.get('strengths', [])),
            "Weaknesses": ", ".join(result.get('weaknesses', [])),
            "Summary": result.get('summary', 'N/A'),
            "Recommendation": result.get('recommendation', 'N/A')
        })

    df = pd.DataFrame(data)
    return df

def generate_pdf_report(results, job_title, categories, total_files, duplicates):
    """Generate PDF Report for Resume Analysis"""

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Title
    pdf.set_font("Arial", "B", 16)
    pdf.cell(200, 10, "AI Resume Shortlisting Report", ln=True, align="C")
    pdf.ln(10)

    # Job Title
    pdf.set_font("Arial", "", 12)
    pdf.cell(200, 10, f"Job Title: {job_title}", ln=True)
    pdf.ln(5)

    # Summary Section
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, "Processing Summary", ln=True)

    pdf.set_font("Arial", "", 11)
    pdf.cell(200, 8, f"Total Files Uploaded: {total_files}", ln=True)
    pdf.cell(200, 8, f"Unique Resumes Analyzed: {len(results)}", ln=True)
    pdf.cell(200, 8, f"Duplicates Skipped: {duplicates}", ln=True)
    pdf.ln(8)

    # Category Breakdown
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, "Results Breakdown", ln=True)

    pdf.set_font("Arial", "", 11)
    for cat, cat_results in categories.items():
        pdf.cell(200, 8, f"{cat}: {len(cat_results)} candidates", ln=True)

    pdf.ln(10)

    # Top Candidates
    pdf.set_font("Arial", "B", 12)
    pdf.cell(200, 10, "Top Candidates", ln=True)

    pdf.set_font("Arial", "", 10)
    for result in results[:10]:
        pdf.multi_cell(
            0,
            7,
            f"Name: {result.get('candidate_name')}\n"
            f"Score: {result.get('fit_score')}%\n"
            f"Role: {result.get('current_role')}\n"
            f"Location: {result.get('location')}\n"
            f"Recommendation: {result.get('recommendation')}\n"
            "----------------------------------------"
        )

    # ✅ Return PDF bytes
    return pdf.output(dest="S").encode("latin-1")

# Sidebar
with st.sidebar:
    st.markdown("""
    <div class="brand-container">
        <div class="brand-icon">🎯</div>
        <div>
            <div class="brand-text">HireSmartAI</div>
            <div class="brand-subtitle">Recruitment OS</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="nav-item active">
        <span>📊</span>
        <span>Dashboard</span>
    </div>
    <div class="nav-item">
        <span>🕐</span>
        <span>History</span>
    </div>
    <div class="nav-item">
        <span>⚙️</span>
        <span>Settings</span>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    st.markdown('<p class="section-label">API Status</p>', unsafe_allow_html=True)
    if GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here":
        st.markdown("""
        <div style="background: #d1fae5; padding: 0.5rem 1rem; border-radius: 8px; display: flex; align-items: center; gap: 8px;">
            <span style="color: #059669;">✓</span>
            <span style="color: #059669; font-size: 0.85rem;">API Key Configured</span>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <div style="background: #fee2e2; padding: 0.5rem 1rem; border-radius: 8px;">
            <span style="color: #dc2626; font-size: 0.85rem;">⚠️ API Key Missing</span>
            <p style="color: #dc2626; font-size: 0.75rem; margin: 0.25rem 0 0 0;">Add to .env file</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown('<p class="section-label">Quick Links</p>', unsafe_allow_html=True)
    st.markdown("[🔗 Get Free API Key](https://console.groq.com)")

    st.markdown("---")
    st.markdown("""
    <div class="user-profile">
        <div class="user-avatar">HR</div>
        <div>
            <div class="user-name">HR Manager</div>
            <div class="user-plan">Pro Plan</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Main content
st.markdown('<h1 class="main-header">New Analysis</h1>', unsafe_allow_html=True)
st.markdown('<p class="main-subtitle">Upload resumes and define job context to find your perfect match.</p>', unsafe_allow_html=True)

# Two column layout for upload and JD
col1, col2 = st.columns(2)

with col1:
    st.markdown('<p class="section-label">Source Files</p>', unsafe_allow_html=True)

    upload_method = st.radio(
        "Upload Method",
        ["📄 Upload Files", "📁 Select Folder"],
        horizontal=True,
        label_visibility="collapsed"
    )

    uploaded_files = []

    if upload_method == "📄 Upload Files":
        uploaded_files = st.file_uploader(
            "Upload Resumes",
            type=['pdf', 'docx', 'doc'],
            accept_multiple_files=True,
            label_visibility="collapsed",
            help="Drag & drop PDF/DOCX files here (supports 250+ files)"
        )

        if uploaded_files:
            st.markdown(f"**{len(uploaded_files)} file(s) selected:**")
            file_chips = ""
            for f in uploaded_files[:3]:
                file_chips += f'<span class="file-chip">📄 {f.name}</span>'
            if len(uploaded_files) > 3:
                file_chips += f'<span class="file-chip">+{len(uploaded_files)-3} others</span>'
            st.markdown(file_chips, unsafe_allow_html=True)

    else:
        # Check if running on Streamlit Cloud
        is_cloud = os.getenv("STREAMLIT_SHARING_MODE") or "streamlit.app" in os.getenv("HOSTNAME", "") or not os.path.exists("C:\\")

        if is_cloud:
            st.warning("⚠️ **Folder selection only works when running locally.** On Streamlit Cloud, please use 'Upload Files' instead.")

        st.markdown("""
        <div style="background: #f9fafb; border: 2px dashed #ddd6fe; border-radius: 12px; padding: 1rem; margin-bottom: 1rem;">
            <p style="color: #6b7280; font-size: 0.9rem; margin: 0;">
                📁 Enter the full path to your resumes folder (Local only)
            </p>
        </div>
        """, unsafe_allow_html=True)

        folder_path_input = st.text_input(
            "Folder Path",
            placeholder=r"C:\Users\HR\Resumes or /home/hr/resumes",
            label_visibility="collapsed",
            key="folder_path_input"
        )

        if folder_path_input:
            # Clean the path - remove quotes, extra spaces, and normalize
            folder_path_clean = folder_path_input.strip().strip('"').strip("'").strip()

            # Handle escaped backslashes (when copied from some sources)
            folder_path_clean = folder_path_clean.replace("\\\\", "\\")

            # Remove any invisible/control characters (keep printable + space)
            folder_path_clean = ''.join(c for c in folder_path_clean if c.isprintable())

            # Handle HTML entities that might come from browser copy
            folder_path_clean = folder_path_clean.replace("&amp;", "&")
            folder_path_clean = folder_path_clean.replace("%20", " ")

            # Remove leading/trailing whitespace again after all cleaning
            folder_path_clean = folder_path_clean.strip()

            # Debug info - show character codes for troubleshooting
            st.caption(f"📂 Path: `{folder_path_clean}` ({len(folder_path_clean)} chars)")

            try:
                folder_path = Path(folder_path_clean)
            except Exception as path_err:
                st.error(f"❌ Invalid path format: {path_err}")
                st.session_state.folder_pdf_paths = []
                folder_path = None

            if folder_path and folder_path.exists() and folder_path.is_dir():
                resume_files = (
                    list(folder_path.glob("*.pdf")) +
                    list(folder_path.glob("*.PDF")) +
                    list(folder_path.glob("*.docx")) +
                    list(folder_path.glob("*.DOCX")) +
                    list(folder_path.glob("*.doc")) +
                    list(folder_path.glob("*.DOC"))
                )

                if resume_files:
                    pdf_count = len([f for f in resume_files if f.suffix.lower() == '.pdf'])
                    docx_count = len([f for f in resume_files if f.suffix.lower() in ['.docx', '.doc']])
                    st.success(f"✅ Found {len(resume_files)} files ({pdf_count} PDF, {docx_count} DOCX/DOC)")
                    st.session_state.folder_pdf_paths = resume_files

                    file_chips = ""
                    for f in resume_files[:3]:
                        file_chips += f'<span class="file-chip">📄 {f.name}</span>'
                    if len(resume_files) > 3:
                        file_chips += f'<span class="file-chip">+{len(resume_files)-3} others</span>'
                    st.markdown(file_chips, unsafe_allow_html=True)

                    include_subfolders = st.checkbox("Include subfolders", value=False)
                    if include_subfolders:
                        resume_files = (
                            list(folder_path.rglob("*.pdf")) +
                            list(folder_path.rglob("*.PDF")) +
                            list(folder_path.rglob("*.docx")) +
                            list(folder_path.rglob("*.DOCX")) +
                            list(folder_path.rglob("*.doc")) +
                            list(folder_path.rglob("*.DOC"))
                        )
                        st.session_state.folder_pdf_paths = resume_files
                        st.info(f"📂 Total with subfolders: {len(resume_files)} files")
                else:
                    st.warning("⚠️ No PDF/DOCX files found in this folder")
                    st.session_state.folder_pdf_paths = []
            elif folder_path:
                # More helpful error message
                if not folder_path.exists():
                    st.error(f"❌ Path does not exist: `{folder_path_clean}`")
                    st.info("💡 Tip: Copy the path from Windows Explorer address bar")
                elif not folder_path.is_dir():
                    st.error(f"❌ Path is not a folder (it's a file): `{folder_path_clean}`")
                else:
                    st.error("❌ Folder not found. Please check the path.")
                st.session_state.folder_pdf_paths = []

with col2:
    st.markdown('<p class="section-label">Job Description</p>', unsafe_allow_html=True)
    job_description = st.text_area(
        "Job Description",
        height=150,
        placeholder="Enter the job description here...\n\nInclude:\n- Role title\n- Required skills\n- Experience requirements\n- Responsibilities",
        label_visibility="collapsed"
    )

    # Nice to Have Skills input (Feature 1)
    st.markdown('<p class="section-label">Nice-to-Have Skills (Optional)</p>', unsafe_allow_html=True)
    nice_to_have_input = st.text_input(
        "Nice-to-Have Skills",
        placeholder="e.g., AWS, Docker, Kubernetes, GraphQL (comma-separated)",
        label_visibility="collapsed"
    )
    nice_to_have_skills = [s.strip() for s in nice_to_have_input.split(",") if s.strip()] if nice_to_have_input else []

    # Feature 6: Suggest Keywords Button (always visible)
    if st.button("🔍 Extract Keywords from JD", use_container_width=True):
        if not job_description:
            st.warning("⚠️ Please enter a job description first")
        elif not GROQ_API_KEY:
            st.warning("⚠️ API key not configured")
        else:
            with st.spinner("Analyzing job description..."):
                client = Groq(api_key=GROQ_API_KEY)
                keywords = extract_keywords_from_jd(client, job_description)
                st.session_state.suggested_keywords = keywords
                st.session_state.job_title = keywords.get('job_title', '')

# Display suggested keywords
if st.session_state.suggested_keywords:
    keywords = st.session_state.suggested_keywords
    st.markdown("---")
    st.markdown("### 📋 Suggested Keywords from JD")

    kw_col1, kw_col2 = st.columns(2)

    with kw_col1:
        if keywords.get('job_title'):
            st.markdown(f"**Job Title:** {keywords['job_title']}")
        if keywords.get('experience_required'):
            st.markdown(f"**Experience Required:** {keywords['experience_required']}")

        if keywords.get('must_have_skills'):
            st.markdown("**Must-Have Skills:**")
            skills_html = "".join([f'<span class="keyword-tag">{s}</span>' for s in keywords['must_have_skills']])
            st.markdown(skills_html, unsafe_allow_html=True)

    with kw_col2:
        if keywords.get('nice_to_have_skills'):
            st.markdown("**Nice-to-Have Skills:**")
            skills_html = "".join([f'<span class="keyword-tag">{s}</span>' for s in keywords['nice_to_have_skills']])
            st.markdown(skills_html, unsafe_allow_html=True)

        if keywords.get('technologies'):
            st.markdown("**Technologies:**")
            tech_html = "".join([f'<span class="keyword-tag">{t}</span>' for t in keywords['technologies']])
            st.markdown(tech_html, unsafe_allow_html=True)

# Analyze button
st.markdown("<br>", unsafe_allow_html=True)
col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
with col_btn2:
    analyze_clicked = st.button("✨ Analyze Candidates", use_container_width=True)

# Process analysis
if analyze_clicked:
    files_to_process = []
    use_folder = upload_method == "📁 Select Folder"

    if use_folder:
        folder_paths = st.session_state.get('folder_pdf_paths', [])
        files_to_process = folder_paths
    else:
        files_to_process = uploaded_files if uploaded_files else []

    # Get job title
    job_title = st.session_state.get('job_title', '')
    if not job_title and st.session_state.suggested_keywords:
        job_title = st.session_state.suggested_keywords.get('job_title', 'Not specified')
    if not job_title:
        job_title = "Not specified"

    # Merge nice-to-have: manual input + extracted JD keywords
    if st.session_state.suggested_keywords:
        extracted_nice = st.session_state.suggested_keywords.get('nice_to_have_skills', [])
        if extracted_nice:
            combined = sorted(set(nice_to_have_skills + extracted_nice), key=str.lower)
            nice_to_have_skills = combined

    if not GROQ_API_KEY or GROQ_API_KEY == "your_groq_api_key_here":
        st.error("⚠️ API key not configured. Please add your Groq API key to the .env file")
    elif not job_description:
        st.error("⚠️ Please enter a job description")
    elif not files_to_process:
        st.error("⚠️ Please upload resumes or select a folder with PDF/DOCX files")
    else:
        try:
            progress_bar = st.progress(0)
            status_text = st.empty()
            time_estimate = st.empty()
            duplicate_info = st.empty()

            total = len(files_to_process)

            # Parallel processing settings
            MAX_WORKERS = min(5, total)  # Limit to 5 parallel requests (Groq rate limit friendly)

            if total > 10:
                est_time = (total // MAX_WORKERS) * 2  # Much faster with parallel
                time_estimate.info(f"⚡ Parallel processing: ~{est_time // 60}m {est_time % 60}s for {total} resumes ({MAX_WORKERS} parallel workers)")

            # Step 1: Extract text from all files first (fast, can be parallelized)
            status_text.text("📄 Extracting text from resumes...")

            extracted_data = []
            seen_hashes = {}
            duplicates_skipped = 0

            for idx, file_item in enumerate(files_to_process):
                if use_folder:
                    file_path = file_item
                    file_name = file_path.name
                    file_ext = file_path.suffix.lower()

                    if file_ext == '.pdf':
                        resume_text = extract_text_from_pdf_path(file_path)
                    elif file_ext in ['.docx', '.doc']:
                        resume_text = extract_text_from_docx_path(file_path)
                    else:
                        resume_text = "Error: Unsupported file format"
                else:
                    file_name = file_item.name
                    file_ext = Path(file_name).suffix.lower()

                    if file_ext == '.pdf':
                        resume_text = extract_text_from_pdf(file_item)
                    elif file_ext in ['.docx', '.doc']:
                        resume_text = extract_text_from_docx(file_item)
                    else:
                        resume_text = "Error: Unsupported file format"

                # Clean the filename (remove platform prefixes, extensions, etc.)
                clean_name = clean_candidate_name(file_name)

                # Try to extract actual name from resume content
                if not resume_text.startswith("Error"):
                    extracted_name = extract_name_from_resume(resume_text)
                    if extracted_name:
                        clean_name = extracted_name

                if resume_text.startswith("Error"):
                    extracted_data.append({
                        "candidate_name": clean_name,
                        "job_title": job_title,
                        "fit_score": 0,
                        "error": resume_text,
                        "verdict": "Error",
                        "current_role": "Unknown",
                        "location": "Not provided",
                        "skills_matched": [],
                        "skills_missing": [],
                        "nice_to_have_matched": [],
                        "summary": "Could not extract text from file",
                        "is_error": True
                    })
                else:
                    content_hash = get_content_hash(resume_text)

                    if content_hash in seen_hashes:
                        duplicates_skipped += 1
                    else:
                        seen_hashes[content_hash] = file_name
                        extracted_data.append({
                            "resume_text": resume_text,
                            "candidate_name": clean_name,
                            "is_error": False
                        })

                progress_bar.progress((idx + 1) / total * 0.3)  # 30% for extraction

            if duplicates_skipped > 0:
                duplicate_info.success(f"✅ Found {duplicates_skipped} duplicate(s) - will be skipped")

            # Step 2: Parallel API calls for analysis
            status_text.text(f"🚀 Analyzing {len([d for d in extracted_data if not d.get('is_error')])} resumes in parallel...")

            results = []
            resumes_to_analyze = [d for d in extracted_data if not d.get('is_error')]
            error_results = [d for d in extracted_data if d.get('is_error')]

            # Add error results directly
            for err in error_results:
                del err['is_error']
                results.append(err)

            # Thread-safe counter for progress
            completed_count = [0]
            lock = threading.Lock()

            def analyze_single_resume(data):
                """Analyze a single resume - called in parallel"""
                client = Groq(api_key=GROQ_API_KEY)  # Each thread gets its own client
                result = analyze_resume(
                    client,
                    data['resume_text'],
                    job_description,
                    nice_to_have_skills,
                    data['candidate_name'],
                    job_title
                )

                with lock:
                    completed_count[0] += 1

                return result

            # Parallel execution
            if resumes_to_analyze:
                with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                    futures = {executor.submit(analyze_single_resume, data): data for data in resumes_to_analyze}

                    for future in as_completed(futures):
                        try:
                            result = future.result()
                            results.append(result)

                            # Update progress (30% extraction + 70% analysis)
                            progress = 0.3 + (completed_count[0] / len(resumes_to_analyze)) * 0.7
                            progress_bar.progress(progress)
                            status_text.text(f"🔍 Analyzed {completed_count[0]}/{len(resumes_to_analyze)} resumes...")

                        except Exception as e:
                            data = futures[future]
                            results.append({
                                "candidate_name": data['candidate_name'],
                                "job_title": job_title,
                                "fit_score": 0,
                                "error": str(e),
                                "verdict": "Error",
                                "current_role": "Unknown",
                                "location": "Not provided",
                                "skills_matched": [],
                                "skills_missing": [],
                                "nice_to_have_matched": [],
                                "summary": f"Error: {str(e)}"
                            })

            progress_bar.progress(1.0)

            results.sort(key=lambda x: x.get('fit_score', 0), reverse=True)

            # Name-based deduplication: if same candidate name appears multiple times,
            # keep only the one with the highest score (using fuzzy matching for robustness)
            seen_names = []
            unique_results = []
            name_duplicates = 0
            for res in results:
                name_key = res.get('candidate_name', '').strip().lower()
                if not name_key or name_key == 'unknown':
                    unique_results.append(res)
                    continue
                # Fuzzy match against already-seen names (handles LLM name variations)
                is_duplicate = False
                for seen in seen_names:
                    if fuzz.token_sort_ratio(name_key, seen) >= 85:
                        is_duplicate = True
                        break
                if is_duplicate:
                    name_duplicates += 1
                else:
                    seen_names.append(name_key)
                    unique_results.append(res)
            results = unique_results
            duplicates_skipped += name_duplicates

            st.session_state.results = results
            st.session_state.analyzed = True
            st.session_state.job_title = job_title

            status_text.empty()
            progress_bar.empty()
            time_estimate.empty()

            st.session_state.duplicates_count = duplicates_skipped
            st.session_state.total_files_processed = total

            if duplicates_skipped > 0:
                duplicate_info.success(f"✅ Processed {len(results)} unique resumes. Skipped {duplicates_skipped} duplicate(s).")
            else:
                duplicate_info.empty()

            st.rerun()

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

# ✅ Display results ONLY after analysis
if st.session_state.analyzed and st.session_state.results:

    results = st.session_state.results
    job_title = st.session_state.get("job_title", "Not specified")
    total_files = st.session_state.get("total_files_processed", len(results))
    duplicates = st.session_state.get("duplicates_count", 0)

    st.markdown("---")

    # Job Title Banner
    st.markdown(f"""
    <div style="background: linear-gradient(135deg, #8b5cf6 0%, #7c3aed 100%);
                padding: 1rem 1.5rem;
                border-radius: 12px;
                margin-bottom: 1rem;">
        <div style="color: white; font-size: 0.8rem;">Analyzing for Position</div>
        <div style="color: white; font-size: 1.25rem; font-weight: 700;">
            {job_title}
        </div>
    </div>
    """, unsafe_allow_html=True)

    # Categorize ALL results
    categories = {
        "Best Fit": [],
        "Strong Fit": [],
        "Average": [],
        "Not a Fit": []
    }

    for res in results:
        verdict = res.get("verdict", "Not a Fit")
        if "Best" in verdict:
            categories["Best Fit"].append(res)
        elif "Strong" in verdict:
            categories["Strong Fit"].append(res)
        elif "Average" in verdict:
            categories["Average"].append(res)
        else:
            categories["Not a Fit"].append(res)

    # ==================== METRIC CARDS ====================
    st.markdown("### 📊 Processing Summary")

    mc1, mc2, mc3, mc4, mc5 = st.columns(5)

    with mc1:
        st.markdown(f"""
        <div style="background: #ede9fe; border-left: 4px solid #7c3aed; padding: 1rem; border-radius: 8px; text-align: center;">
            <div style="font-size: 2rem; font-weight: 700; color: #7c3aed;">{len(categories["Best Fit"])}</div>
            <div style="font-size: 0.85rem; color: #5b21b6;">Best Fit</div>
        </div>
        """, unsafe_allow_html=True)

    with mc2:
        st.markdown(f"""
        <div style="background: #d1fae5; border-left: 4px solid #059669; padding: 1rem; border-radius: 8px; text-align: center;">
            <div style="font-size: 2rem; font-weight: 700; color: #059669;">{len(categories["Strong Fit"])}</div>
            <div style="font-size: 0.85rem; color: #047857;">Strong Fit</div>
        </div>
        """, unsafe_allow_html=True)

    with mc3:
        st.markdown(f"""
        <div style="background: #fef3c7; border-left: 4px solid #d97706; padding: 1rem; border-radius: 8px; text-align: center;">
            <div style="font-size: 2rem; font-weight: 700; color: #d97706;">{len(categories["Average"])}</div>
            <div style="font-size: 0.85rem; color: #b45309;">Average</div>
        </div>
        """, unsafe_allow_html=True)

    with mc4:
        st.markdown(f"""
        <div style="background: #fee2e2; border-left: 4px solid #dc2626; padding: 1rem; border-radius: 8px; text-align: center;">
            <div style="font-size: 2rem; font-weight: 700; color: #dc2626;">{len(categories["Not a Fit"])}</div>
            <div style="font-size: 0.85rem; color: #b91c1c;">Not a Fit</div>
        </div>
        """, unsafe_allow_html=True)

    with mc5:
        st.markdown(f"""
        <div style="background: #e0e7ff; border-left: 4px solid #4f46e5; padding: 1rem; border-radius: 8px; text-align: center;">
            <div style="font-size: 2rem; font-weight: 700; color: #4f46e5;">{duplicates}</div>
            <div style="font-size: 0.85rem; color: #3730a3;">Duplicates Skipped</div>
        </div>
        """, unsafe_allow_html=True)

    st.caption(f"Total uploaded: {total_files} | Unique analyzed: {len(results)} | Duplicates skipped: {duplicates}")

    st.markdown("---")

    # ==================== CATEGORY SECTIONS WITH VIEW DETAILS ====================
    category_icons = {"Best Fit": "🏆", "Strong Fit": "✅", "Average": "📋", "Not a Fit": "❌"}
    category_colors = {"Best Fit": "#7c3aed", "Strong Fit": "#059669", "Average": "#d97706", "Not a Fit": "#dc2626"}

    for category, category_results in categories.items():
        if not category_results:
            continue

        color = category_colors[category]
        icon = category_icons[category]

        st.markdown(f"""
        <div style="display: flex; align-items: center; gap: 10px; margin: 1.5rem 0 0.75rem 0; padding-bottom: 0.5rem; border-bottom: 3px solid {color};">
            <span style="font-size: 1.5rem;">{icon}</span>
            <span style="font-size: 1.25rem; font-weight: 700; color: #1f2937;">{category}</span>
            <span style="background: {color}; color: white; padding: 0.2rem 0.6rem; border-radius: 12px; font-size: 0.85rem; font-weight: 600;">{len(category_results)}</span>
        </div>
        """, unsafe_allow_html=True)

        for idx, candidate in enumerate(category_results):
            cand_name = candidate.get("candidate_name", "Unknown")
            cand_score = candidate.get("fit_score", 0)
            cand_role = candidate.get("current_role", "N/A")
            cand_loc = candidate.get("location", "N/A")
            cand_exp = candidate.get("experience_years", 0)
            cand_email = candidate.get("email", "N/A")
            cand_phone = candidate.get("phone", "N/A")
            cand_rec = candidate.get("recommendation", "N/A")

            # Candidate header
            st.markdown(f"""
            **👤 {cand_name}** — Score: **{cand_score}%** | 📍 {cand_loc} | 💼 {cand_role} | 📅 {cand_exp} yrs | 🎯 {cand_rec}
            """)

            # Expandable View Details
            with st.expander(f"View Details - {cand_name}"):

                # Score Breakdown
                breakdown = candidate.get("score_breakdown", {})
                st.markdown("**📊 Score Breakdown:**")
                bd_c1, bd_c2, bd_c3, bd_c4, bd_c5 = st.columns(5)
                bd_c1.metric("Skills", f"{breakdown.get('skills_score', 0)}/40")
                bd_c2.metric("Experience", f"{breakdown.get('experience_score', 0)}/25")
                bd_c3.metric("Nice-to-Have", f"{breakdown.get('nice_to_have_score', 0)}/15")
                bd_c4.metric("Education", f"{breakdown.get('education_score', 0)}/10")
                bd_c5.metric("Relevance", f"{breakdown.get('relevance_score', 0)}/10")

                st.markdown("---")

                det_c1, det_c2 = st.columns(2)

                with det_c1:
                    # Contact Info
                    st.markdown(f"**📧 Email:** {cand_email}")
                    st.markdown(f"**📱 Phone:** {cand_phone}")
                    st.markdown(f"**🎓 Education:** {candidate.get('education_level', 'N/A')}")

                    # Skills Matched
                    matched = candidate.get("skills_matched", [])
                    if matched:
                        st.markdown("**✅ Skills Matched:**")
                        skills_html = "".join([f'<span class="skill-matched">{s}</span>' for s in matched])
                        st.markdown(skills_html, unsafe_allow_html=True)

                    # Skills Missing
                    missing = candidate.get("skills_missing", [])
                    if missing:
                        st.markdown("**❌ Skills Missing:**")
                        missing_html = "".join([f'<span class="skill-missing">{s}</span>' for s in missing])
                        st.markdown(missing_html, unsafe_allow_html=True)

                    # Nice-to-Have Matched
                    nice = candidate.get("nice_to_have_matched", [])
                    if nice:
                        st.markdown("**🎁 Nice-to-Have Matched:**")
                        nice_html = "".join([f'<span class="skill-nice">{s}</span>' for s in nice])
                        st.markdown(nice_html, unsafe_allow_html=True)

                with det_c2:
                    # Strengths
                    strengths = candidate.get("strengths", [])
                    if strengths:
                        st.markdown("**💪 Strengths:**")
                        for s in strengths:
                            st.markdown(f"- {s}")

                    # Weaknesses
                    weaknesses = candidate.get("weaknesses", [])
                    if weaknesses:
                        st.markdown("**⚠️ Weaknesses:**")
                        for w in weaknesses:
                            st.markdown(f"- {w}")

                    # Summary
                    summary = candidate.get("summary", "")
                    if summary:
                        st.markdown(f"**📝 Summary:** {summary}")

                    # TF-IDF Similarity
                    tfidf = candidate.get("tfidf_similarity", 0)
                    if tfidf:
                        st.markdown(f"**📐 Resume-JD Similarity:** {tfidf}%")

            st.markdown("---")

    # ==================== DOWNLOAD REPORTS ====================
    st.markdown("## 📥 Download Reports")

    download_col1, download_col2, download_col3 = st.columns(3)

    df = create_excel_report(results, categories, total_files, duplicates, job_title)

    with download_col1:
        st.download_button(
            "📄 Download CSV",
            df.to_csv(index=False),
            "resume_analysis_report.csv",
            "text/csv"
        )

    with download_col2:
        excel_buffer = io.BytesIO()
        with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
            df.to_excel(writer, index=False)
        excel_buffer.seek(0)
        st.download_button(
            "📊 Download Excel",
            excel_buffer,
            "resume_analysis_report.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    with download_col3:
        pdf_bytes = generate_pdf_report(results, job_title, categories, total_files, duplicates)
        st.download_button(
            "📑 Download PDF",
            pdf_bytes,
            "resume_analysis_report.pdf",
            "application/pdf"
        )


# ✅ Footer ALWAYS outside block
st.markdown("""
<div style="text-align: center; padding: 2rem;
            color: #9ca3af; font-size: 0.85rem;">
    Built for Mamsys Hackathon 2025 | Powered by Groq AI
</div>
""", unsafe_allow_html=True)
